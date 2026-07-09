"""Build the ResearchSense FYP report (Chapters 1 to 4) as a .docx file.

Mirrors the Bahria University LaTeX thesis template: a title page, an abstract,
then four numbered chapters with sections. Run from this folder:

    python build_report.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import report_content as c

OUTPUT = Path(__file__).resolve().parents[3] / "ResearchSense_FYP_Report.docx"
BLACK = RGBColor(0, 0, 0)
SERIF = "Times New Roman"


def style_document(doc: Document) -> None:
    """Set a plain black serif look similar to the LaTeX thesis."""
    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.4
    for level, size in ((1, 22), (2, 15)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = SERIF
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLACK


def centered(doc: Document, text: str, size: int, bold: bool = False,
             space_before: int = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = SERIF
    run.font.size = Pt(size)
    run.bold = bold


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def title_page(doc: Document) -> None:
    centered(doc, c.UNIVERSITY, 16, bold=True, space_before=36)
    centered(doc, "PAKISTAN", 11)
    centered(doc, c.AUTHOR, 14, bold=True, space_before=90)
    centered(doc, c.ENROLLMENT, 12)
    centered(doc, c.TITLE, 20, bold=True, space_before=60)
    centered(doc, c.DEGREE, 13, bold=True, space_before=90)
    centered(doc, f"Supervisor: {c.SUPERVISOR}", 12, space_before=40)
    centered(doc, f"Co-Supervisor: {c.CO_SUPERVISOR}", 12)
    centered(doc, c.DEPARTMENT, 12, space_before=40)
    centered(doc, c.UNIVERSITY, 12)
    centered(doc, c.DATE, 12, space_before=20)
    doc.add_page_break()


def abstract_page(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    for para in c.ABSTRACT:
        body(doc, para)
    doc.add_page_break()


def render_section(doc: Document, section: tuple) -> None:
    """A section is (title, paragraphs) or (title, paragraphs, bullet_items)."""
    heading, paragraphs = section[0], section[1]
    doc.add_heading(heading, level=2)
    for para in paragraphs:
        body(doc, para)
    if len(section) == 3:
        bullets(doc, section[2])


def data_dictionary(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Table", "Key Fields", "Purpose")
    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for name, fields, purpose in c.DB_TABLE:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = fields
        row[2].text = purpose
    doc.add_paragraph()


def render_chapter(doc: Document, number: int, chapter: dict,
                   *, db_after: str | None = None,
                   refs: list[str] | None = None) -> None:
    doc.add_heading(f"Chapter {number}", level=1)
    doc.add_heading(chapter["title"], level=1)
    for section in chapter["sections"]:
        render_section(doc, section)
        if db_after and section[0] == db_after:
            data_dictionary(doc)
    if refs:
        doc.add_heading("References", level=2)
        for i, ref in enumerate(refs, start=1):
            doc.add_paragraph(f"[{i}] {ref}")
    doc.add_page_break()


def main() -> None:
    doc = Document()
    style_document(doc)

    title_page(doc)
    abstract_page(doc)
    render_chapter(doc, 1, c.CH1)
    render_chapter(doc, 2, c.CH2, refs=c.REFERENCES)
    render_chapter(doc, 3, c.CH3)
    render_chapter(doc, 4, c.CH4, db_after="Database Design")

    doc.save(OUTPUT)
    paragraphs = len(doc.paragraphs)
    print(f"Saved {OUTPUT}")
    print(f"Paragraphs: {paragraphs}")


if __name__ == "__main__":
    main()
